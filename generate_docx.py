"""Generate PROJECT_DOCUMENTATION.docx from the project documentation content."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Helper functions ---

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table


# ===== TITLE PAGE =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\nGen Z Loyalty Analysis\nAmerican Airlines')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('\nBUS 480 Project Documentation')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ===== TABLE OF CONTENTS =====
add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Data Overview',
    '3. Data Processing and Cleaning',
    '4. Data Overview and Summary After Cleaning',
    '5. Exploratory Data Analysis (EDA)',
    '6. Machine Learning Methods — Problem 1',
    '7. Machine Learning Methods — Problem 2',
    '8. Results and Insights',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ===== 1. EXECUTIVE SUMMARY =====
add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'This project investigates two core business questions for American Airlines:'
)

add_para('Problem 1 — Who is Gen Z compared to other generations?', bold=True)
doc.add_paragraph(
    'How does Gen Z differ from Millennials and Gen X in loyalty behavior? '
    'We answer this through exploratory data analysis (EDA) and K-Means clustering, '
    'comparing loyalty enrollment rates, engagement levels, and booking patterns across generational cohorts.'
)

add_para('Problem 2 — What motivates Gen Z loyalty?', bold=True)
doc.add_paragraph(
    'What actually makes Gen Z stay loyal? We answer this using two gradient boosting models '
    '(XGBoost and LightGBM) paired with SHAP explainability to identify and rank the specific '
    'service factors that drive Gen Z loyalty enrollment.'
)

doc.add_paragraph(
    'The analysis uses an airline passenger satisfaction dataset with 129,880 total records across '
    '25 columns, covering demographic information, 14 in-flight service ratings, flight details, '
    'and satisfaction outcomes.'
)

doc.add_page_break()

# ===== 2. DATA OVERVIEW =====
add_heading('2. Data Overview', level=1)

add_heading('2.1 Data Sources', level=2)

add_table(
    ['Dataset', 'Rows', 'Columns', 'File Size'],
    [
        ['train.csv', '103,904', '25', '12 MB'],
        ['test.csv', '25,976', '25', '2.9 MB'],
        ['Combined', '129,880', '25', '~15 MB'],
    ]
)

doc.add_paragraph(
    '\nBoth datasets share identical column structures and were combined for the full analysis.'
)

add_heading('2.2 Column Descriptions', level=2)

add_table(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'Integer', 'Unique passenger identifier'],
        ['Gender', 'Categorical', 'Male or Female'],
        ['Customer Type', 'Categorical', '"Loyal Customer" or "disloyal Customer"'],
        ['Age', 'Integer', 'Passenger age (7–85 range)'],
        ['Type of Travel', 'Categorical', '"Business travel" or "Personal Travel"'],
        ['Class', 'Categorical', '"Business", "Eco", or "Eco Plus"'],
        ['Flight Distance', 'Integer', 'Distance of the flight in miles'],
        ['Departure Delay in Minutes', 'Float', 'Minutes delayed at departure'],
        ['Arrival Delay in Minutes', 'Float', 'Minutes delayed at arrival'],
        ['satisfaction', 'Categorical', 'Target: "satisfied" or "neutral or dissatisfied"'],
    ]
)

add_heading('14 Service Rating Columns (rated 0–5 scale)', level=3)

add_table(
    ['Service Dimension', 'What It Measures'],
    [
        ['Inflight wifi service', 'Quality of in-flight Wi-Fi connectivity'],
        ['Departure/Arrival time convenient', 'Convenience of scheduled flight times'],
        ['Ease of Online booking', 'Ease of completing an online booking'],
        ['Gate location', 'Convenience of gate assignment'],
        ['Food and drink', 'Quality of in-flight food and beverages'],
        ['Online boarding', 'Smoothness of online boarding process'],
        ['Seat comfort', 'Physical comfort of seating'],
        ['Inflight entertainment', 'Quality of entertainment options'],
        ['On-board service', 'Quality of crew service during flight'],
        ['Leg room service', 'Adequacy of legroom space'],
        ['Baggage handling', 'Efficiency of baggage handling'],
        ['Checkin service', 'Quality of check-in experience'],
        ['Inflight service', 'Overall in-flight service quality'],
        ['Cleanliness', 'Cleanliness of the aircraft'],
    ]
)

doc.add_paragraph(
    '\nRating scale: 0 = no service / not applicable, 1 = very dissatisfied, 5 = very satisfied.'
)

doc.add_page_break()

# ===== 3. DATA PROCESSING AND CLEANING =====
add_heading('3. Data Processing and Cleaning', level=1)

add_heading('3.1 Missing Value Treatment', level=2)
doc.add_paragraph(
    'Only one column had missing values:'
)
add_table(
    ['Column', 'Missing Count', 'Imputation Method'],
    [['Arrival Delay in Minutes', '~310 rows (in train set)', 'Median imputation']]
)
doc.add_paragraph(
    '\nMedian was chosen over mean because delay distributions are heavily right-skewed '
    '(most flights have 0 delay, a few have very large delays). Median is robust to these outliers.'
)

add_heading('3.2 Column Removal', level=2)
doc.add_paragraph(
    'The unnamed index column and id column were dropped as they carry no analytical value.'
)

add_heading('3.3 Feature Engineering', level=2)
doc.add_paragraph('Three new features were created from existing columns:')

add_para('Generation Labels', bold=True)
doc.add_paragraph(
    'Derived from Age using standard generational cutoffs (2024 reference year):'
)
add_table(
    ['Generation', 'Age Range', 'Birth Year Range'],
    [
        ['Gen Z', '≤ 27', '~1997–2012'],
        ['Millennial', '28–43', '~1981–1996'],
        ['Gen X', '44–59', '~1965–1980'],
        ['Boomer', '60+', 'Before 1965'],
    ]
)

add_para('\nEngagement Score', bold=True)
doc.add_paragraph(
    'The mean of all 14 service rating columns for each passenger. This serves as a composite '
    'proxy for how engaged and satisfied a member is with the airline\'s service touchpoints.'
)

add_para('Binary Flags', bold=True)
add_bullet(' 1 if Customer Type == "Loyal Customer", 0 otherwise. This is the target variable for Problem 2.', bold_prefix='is_loyal:')
add_bullet(' 1 if satisfaction == "satisfied", 0 otherwise.', bold_prefix='is_satisfied:')

add_heading('3.4 Categorical Encoding (for ML Models)', level=2)
doc.add_paragraph(
    'For the gradient boosting models in Problem 2, three categorical columns were label-encoded:'
)
add_bullet('Class → class_encoded')
add_bullet('Type of Travel → travel_type_encoded')
add_bullet('Gender → gender_encoded')
doc.add_paragraph(
    'Label encoding is appropriate here because XGBoost and LightGBM handle encoded categoricals '
    'natively via their tree-splitting mechanism. No one-hot encoding is needed.'
)

doc.add_page_break()

# ===== 4. DATA OVERVIEW AFTER CLEANING =====
add_heading('4. Data Overview and Summary After Cleaning', level=1)

add_heading('4.1 Final Dataset Shape', level=2)
doc.add_paragraph(
    'After combining train and test sets, imputing missing values, and engineering features, '
    'the final dataset contains:'
)
add_bullet('129,880 rows (passengers)')
add_bullet('0 missing values')
add_bullet('25 original columns + 5 engineered columns (Generation, is_loyal, engagement_score, is_satisfied, cluster)')

add_heading('4.2 Target Variable Distribution', level=2)
add_bullet(' Binary split between "Loyal Customer" and "disloyal Customer". '
           'Stratified train-test splits are used to preserve the class ratio.', bold_prefix='is_loyal (Problem 2 target):')
add_bullet(' Approximately balanced between "satisfied" and "neutral or dissatisfied."', bold_prefix='satisfaction (original target):')

add_heading('4.3 Feature Summary', level=2)
add_table(
    ['Feature Category', 'Count', 'Range / Scale'],
    [
        ['Service ratings', '14', '0–5 integer scale'],
        ['Delay features', '2', '0–1600+ minutes (continuous)'],
        ['Flight distance', '1', '31–4983 miles (continuous)'],
        ['Age', '1', '7–85 years (continuous)'],
        ['Encoded categoricals', '3', 'Label-encoded integers'],
        ['Total model features', '21', ''],
    ]
)

doc.add_page_break()

# ===== 5. EDA =====
add_heading('5. Exploratory Data Analysis (EDA)', level=1)

add_heading('5.1 Purpose', level=2)
doc.add_paragraph(
    'EDA directly answers Problem 1 — the generational comparisons are descriptive analysis. '
    'No ML model is needed. The goal is to summarize, visualize, and statistically test whether '
    'Gen Z differs meaningfully from other generations across loyalty, engagement, and service preferences.'
)

add_heading('5.2 Analyses Performed', level=2)

add_para('Loyalty Enrollment Rate by Generation', bold=True)
doc.add_paragraph(
    'Bar chart showing the percentage of each generation enrolled as "Loyal Customer." '
    'This is the most direct answer to "how does Gen Z differ in loyalty behavior?"'
)

add_para('Engagement Score Distribution by Generation', bold=True)
doc.add_paragraph(
    'Boxplots and violin plots showing the spread of the composite engagement score '
    '(mean service rating) by generation. This reveals whether Gen Z rates services lower overall, '
    'has more variance, or clusters at certain levels.'
)

add_para('Satisfaction Rate by Generation', bold=True)
doc.add_paragraph('Bar chart comparing the proportion of "satisfied" passengers across generations.')

add_para('Generation Comparison Table', bold=True)
doc.add_paragraph(
    'A single summary table aggregating count, loyalty rate, satisfaction rate, mean/median engagement, '
    'average flight distance, and average delays for each generation. This is the key deliverable table for Problem 1.'
)

add_para('Service Ratings Heatmap by Generation', bold=True)
doc.add_paragraph(
    'A 4-row (generation) × 14-column (service) heatmap showing mean ratings. This identifies which '
    'specific services Gen Z rates differently — for example, if Gen Z rates Wi-Fi higher but food lower, '
    'that informs targeted investment.'
)

add_para('Travel Class and Type Distribution', bold=True)
doc.add_paragraph(
    'Stacked bar charts showing how each generation distributes across travel classes (Business, Eco, Eco Plus) '
    'and travel types (Business, Personal). This contextualizes loyalty differences — if Gen Z flies economy more often, '
    'their lower engagement may reflect class-of-service effects rather than generational preferences.'
)

add_para('Correlation Heatmap', bold=True)
doc.add_paragraph(
    'A triangular heatmap of all 14 service ratings plus flight distance, delays, loyalty, satisfaction, and '
    'engagement score. This reveals multicollinearity between features and identifies which service dimensions '
    'correlate most strongly with loyalty and satisfaction.'
)

add_heading('5.3 Statistical Tests', level=2)
doc.add_paragraph(
    'Three hypothesis tests confirm whether observed generational differences are statistically significant '
    'or could be due to chance:'
)
add_table(
    ['Test', 'Null Hypothesis', 'Method'],
    [
        ['Loyalty vs Generation', 'Loyalty enrollment is independent of generation', 'Chi-square test of independence'],
        ['Engagement vs Generation', 'Mean engagement score is equal across generations', 'One-way ANOVA (F-test)'],
        ['Satisfaction vs Generation', 'Satisfaction is independent of generation', 'Chi-square test of independence'],
    ]
)
doc.add_paragraph('\nAll tests use alpha = 0.05. If p-value < 0.05, the generational difference is statistically significant.')

add_heading('5.4 Tools Used', level=2)
doc.add_paragraph(
    'pandas for data aggregation and summary statistics, matplotlib and seaborn for visualization, '
    'scipy.stats for chi-square and ANOVA tests.'
)

doc.add_page_break()

# ===== 6. ML METHODS — PROBLEM 1 =====
add_heading('6. Machine Learning Methods — Problem 1', level=1)

add_heading('6.1 K-Means Clustering', level=2)

add_heading('Introduction', level=3)
doc.add_paragraph(
    'K-Means is an unsupervised machine learning algorithm that groups data points into K clusters '
    'based on behavioral similarity. It works by:'
)
doc.add_paragraph('1. Randomly initializing K cluster centers', style='List Number')
doc.add_paragraph('2. Assigning each data point to its nearest center (Euclidean distance)', style='List Number')
doc.add_paragraph('3. Recomputing centers as the mean of assigned points', style='List Number')
doc.add_paragraph('4. Repeating steps 2–3 until convergence', style='List Number')
doc.add_paragraph('No labels are needed — the algorithm discovers natural structure in the data on its own.')

add_heading('Why K-Means Is Important for This Problem', level=3)
doc.add_paragraph(
    'Generation is a demographic label, not a behavioral identity. Two Gen Z passengers can behave '
    'completely differently — one may be a high-engagement business traveler, another a price-sensitive '
    'occasional flyer. K-Means finds the actual behavioral segments that cut across generational lines.'
)
doc.add_paragraph(
    'By examining which generation dominates each cluster, we learn whether generational labels align '
    'with behavioral reality or whether demographic assumptions are misleading. This is critical for '
    'American Airlines because marketing campaigns should target behavioral personas, not just age groups.'
)

add_heading('Implementation Details', level=3)
add_bullet('Input features (18 total): All 14 service ratings + Flight Distance + Departure Delay + Arrival Delay + Age')
add_bullet('Preprocessing: StandardScaler standardization (zero mean, unit variance). Critical because K-Means uses Euclidean distance.')
add_bullet('Cluster selection: Elbow method (inertia vs K) and silhouette score, evaluating K = 2 to 8.')
add_bullet('Final model: K=4 clusters, fit on the full standardized dataset (129,880 passengers).')

add_heading('Results and Output', level=3)
doc.add_paragraph('The K-Means model produces:')
add_bullet('Elbow chart: Visual confirmation that K=4 captures most variance without over-segmenting')
add_bullet('Silhouette scores: Quantitative validation of cluster quality for each K')
add_bullet('Cluster behavioral profiles: Mean values of all features within each cluster, revealing distinct personas')
add_bullet('Cluster center heatmap: Standardized feature values per cluster')
add_bullet('Generation composition: What percentage of each cluster belongs to each generation, and vice versa')
add_bullet('Persona summary table: Loyalty rate, engagement score, satisfaction rate, and size per cluster')

doc.add_paragraph(
    '\nExpected outcome: We expect 3–5 distinct behavioral segments. Gen Z may concentrate in clusters '
    'characterized by lower engagement or different service priorities, but the clusters should reveal that '
    'behavioral differences within a generation can be larger than differences between generations.'
)

doc.add_page_break()

# ===== 7. ML METHODS — PROBLEM 2 =====
add_heading('7. Machine Learning Methods — Problem 2', level=1)

add_heading('7.1 XGBoost (Extreme Gradient Boosting)', level=2)

add_heading('Introduction', level=3)
doc.add_paragraph(
    'XGBoost is a supervised gradient boosting algorithm that builds an ensemble of decision trees '
    'sequentially. Each tree corrects the errors of the ones before it — this error-correction chain '
    'is the "boosting" mechanism. The process works step by step:'
)
doc.add_paragraph('1. Start with a baseline prediction (mean loyalty rate for all Gen Z members)', style='List Number')
doc.add_paragraph('2. Calculate residuals — how wrong was each prediction?', style='List Number')
doc.add_paragraph('3. Grow a decision tree on those residuals — which features best explain the errors?', style='List Number')
doc.add_paragraph('4. Multiply the tree\'s output by a small learning rate and add it to the running prediction', style='List Number')
doc.add_paragraph('5. Repeat 100–300 times, each tree attacking remaining errors', style='List Number')
doc.add_paragraph('6. Early stopping halts training when validation loss stops improving', style='List Number')
doc.add_paragraph('7. Extract feature importance — Gain (how much each feature reduced error) ranks the loyalty drivers', style='List Number')

add_heading('Why XGBoost Is Important for This Problem', level=3)
doc.add_paragraph(
    'The relationship between loyalty drivers and outcomes is non-linear. A small improvement in Wi-Fi quality '
    'may do nothing until a threshold is crossed, then drive a large enrollment jump. XGBoost handles this '
    'natively through its tree-based architecture.'
)
doc.add_paragraph(
    'It also handles mixed data types (Likert-scale survey scores, continuous distances, binary flags) '
    'without manual feature engineering. The built-in L1 and L2 regularization prevents overfitting, '
    'which is important when the Gen Z subset may be smaller than the full dataset.'
)
doc.add_paragraph(
    'Most importantly, XGBoost produces feature importance rankings that directly answer '
    '"what makes Gen Z stay loyal?" — both enrollment probability AND driver ranking in one model.'
)

add_heading('Implementation Details', level=3)
add_bullet('Target variable: is_loyal (binary: 1 = Loyal Customer, 0 = disloyal Customer)')
add_bullet('Feature set (21 features): 14 service ratings + Flight Distance + Delays + Age + 3 encoded categoricals')
add_bullet('Data split: Gen Z subset only, 80% train / 20% test, stratified by target variable')
add_bullet('300 trees with early stopping at 20 rounds of no improvement')
add_bullet('Max depth 5, learning rate 0.1, subsample 0.8, colsample_bytree 0.8')
add_bullet('L1 regularization (reg_alpha) = 1.0 and L2 regularization (reg_lambda) = 1.0')
add_bullet('Validation: 5-fold stratified cross-validation on the full Gen Z dataset')

add_heading('Results and Output', level=3)
add_bullet('ROC-AUC score: Measures ability to distinguish loyal from non-loyal Gen Z members (0.5 = random, 1.0 = perfect)')
add_bullet('Classification report: Precision, recall, and F1-score for both classes')
add_bullet('Cross-validation ROC-AUC: Mean and standard deviation across 5 folds')
add_bullet('Feature importance chart (Gain): Ranking all 21 features by contribution to reducing prediction error')

doc.add_paragraph(
    '\nExpectation: We expect service-quality features (Wi-Fi, online boarding, seat comfort) to rank higher '
    'than operational features (delays, flight distance) as loyalty drivers, since service experience is more '
    'closely tied to brand loyalty decisions.'
)

add_heading('7.2 LightGBM (Light Gradient Boosting Machine)', level=2)

add_heading('Introduction', level=3)
doc.add_paragraph(
    'LightGBM is Microsoft\'s implementation of gradient boosting. It uses the same core concept as '
    'XGBoost — sequential trees correcting each other\'s errors — but differs in two key ways:'
)
add_bullet(' Instead of growing trees level-by-level, LightGBM finds the single leaf that '
           'reduces error the most and splits that one first, regardless of tree balance. '
           'This produces more accurate trees with fewer splits.', bold_prefix='Leaf-wise growth:')
add_bullet(' Continuous features are binned into discrete histograms before splitting, '
           'dramatically reducing computation time (typically 3–10x faster than XGBoost).', bold_prefix='Histogram-based binning:')

add_heading('Why Running Both Models Is Important', level=3)
doc.add_paragraph('Running two models is standard ML practice for validation:')
add_bullet(' If both XGBoost and LightGBM rank the same feature as the #1 loyalty driver, '
           'that finding is robust and can be reported with high confidence.', bold_prefix='Model agreement = confidence.')
add_bullet(' If XGBoost flags a feature as top-3 but LightGBM doesn\'t, it usually reveals '
           'a segment-specific effect worth investigating.', bold_prefix='Model disagreement = insight.')
doc.add_paragraph(
    'The Spearman rank correlation between the two models\' feature rankings quantifies this agreement '
    '(> 0.7 = strong agreement).'
)

add_heading('Results and Output', level=3)
doc.add_paragraph(
    'Same metrics as XGBoost: ROC-AUC, classification report, cross-validation scores, and feature importance chart. '
    'Additionally, a model comparison table presents both models side-by-side, and a feature ranking comparison '
    'table flags agreements and divergences between the two models.'
)

add_heading('7.3 SHAP Values (SHapley Additive exPlanations)', level=2)

add_heading('Introduction', level=3)
doc.add_paragraph(
    'SHAP is rooted in cooperative game theory (Shapley values). Standard feature importance tells you '
    'what mattered globally across all predictions. SHAP tells you how much each feature contributed to '
    'a specific member\'s prediction and in which direction.'
)
doc.add_paragraph(
    'For example, standard importance might say "Online boarding is important." SHAP says "for this specific '
    'Gen Z member, their low online boarding score of 2 reduced their loyalty probability by 12 percentage points, '
    'while their high seat comfort score of 5 increased it by 8 points."'
)

add_heading('Why SHAP Is Important for This Problem', level=3)
doc.add_paragraph('SHAP transforms model outputs into actionable business recommendations:')
add_bullet(' Invest in that service area', bold_prefix='High positive SHAP on a service feature →')
add_bullet(' Address the pain point', bold_prefix='High negative SHAP on a feature →')
add_bullet(' Segment campaigns rather than one-size-fits-all', bold_prefix='High SHAP variance on a feature →')
add_bullet(' Pair those improvements together', bold_prefix='SHAP interactions between features →')

add_heading('Results and Output', level=3)
doc.add_paragraph('SHAP produces four types of visualizations:')
add_bullet(' Every dot is one Gen Z member\'s SHAP value for one feature. '
           'Color shows feature value (red = high, blue = low). Position shows impact on prediction. '
           'This is the most information-dense visualization in the analysis.', bold_prefix='1. SHAP Summary Plot (Beeswarm):')
add_bullet(' Mean absolute SHAP values per feature — a cleaned-up global ranking.', bold_prefix='2. SHAP Bar Plot:')
add_bullet(' Scatter plots showing feature value vs SHAP contribution for the top 4 features. '
           'These reveal non-linear thresholds.', bold_prefix='3. SHAP Dependence Plots:')
add_bullet(' Individual member explanations showing exactly how each feature pushed the prediction '
           'from the base rate to the final probability. Two contrasting examples are shown — '
           'one loyal and one not loyal.', bold_prefix='4. SHAP Waterfall Plots:')

doc.add_page_break()

# ===== 8. RESULTS AND INSIGHTS =====
add_heading('8. Results and Insights', level=1)

add_heading('8.1 Problem 1 — Who Is Gen Z?', level=2)

add_para('Loyalty Enrollment', bold=True)
doc.add_paragraph(
    'The generational comparison table and bar charts reveal how Gen Z\'s loyalty enrollment rate '
    'compares to Millennials, Gen X, and Boomers. Statistical testing (chi-square) confirms whether '
    'this difference is significant.'
)

add_para('Engagement Levels', bold=True)
doc.add_paragraph(
    'Boxplots and violin plots show the distribution of the composite engagement score by generation. '
    'Key observations include whether Gen Z has a lower median engagement, wider spread (more heterogeneous), '
    'or a bimodal distribution suggesting two distinct sub-segments.'
)

add_para('Service Preferences', bold=True)
doc.add_paragraph(
    'The generation × service heatmap identifies which of the 14 service dimensions Gen Z rates differently. '
    'This directly informs where American Airlines should invest to improve Gen Z satisfaction.'
)

add_para('Travel Behavior', bold=True)
doc.add_paragraph(
    'Class and travel type distributions contextualize loyalty differences — Gen Z\'s lower loyalty may partly '
    'reflect that they fly economy and for personal travel more often, not necessarily that they dislike the airline.'
)

add_para('Behavioral Clusters', bold=True)
doc.add_paragraph(
    'K-Means identifies 4 behavioral personas that cut across generational lines. The generation composition '
    'of each cluster reveals whether Gen Z concentrates in specific behavioral segments or distributes evenly. '
    'This is critical because it determines whether generational targeting or behavioral targeting is more '
    'effective for marketing.'
)

add_heading('8.2 Problem 2 — What Motivates Gen Z Loyalty?', level=2)

add_para('Model Performance', bold=True)
doc.add_paragraph(
    'Both XGBoost and LightGBM are evaluated on ROC-AUC, accuracy, F1-score, and 5-fold cross-validation. '
    'The performance comparison table shows whether the models achieve reliable predictive power on the Gen Z subset.'
)

add_para('Top Loyalty Drivers', bold=True)
doc.add_paragraph(
    'The feature importance rankings from both models are compared. Features where both models agree on high '
    'importance are the most confident loyalty drivers. The Spearman rank correlation quantifies overall model agreement.'
)

add_para('SHAP Insights', bold=True)
doc.add_paragraph(
    'The SHAP summary plot reveals not just which features matter, but the direction and magnitude of their impact. '
    'The dependence plots expose non-linear relationships and threshold effects. The waterfall plots provide concrete, '
    'member-level examples that illustrate how the drivers work in practice.'
)

add_para('Actionable Outputs', bold=True)
add_bullet('A ranked list of Gen Z loyalty drivers (agreed upon by both models)')
add_bullet('Enrollment probability scores for every Gen Z member in the test set')
add_bullet('Specific service areas where investment would most impact Gen Z loyalty')
add_bullet('Evidence for whether one-size-fits-all or segmented campaigns are appropriate (based on SHAP variance)')

add_heading('8.3 Method Justification Summary', level=2)

add_table(
    ['Method', 'Problem', 'Why Chosen'],
    [
        ['EDA', '1', 'Descriptive comparisons are the direct answer — no model needed for "how do generations differ"'],
        ['K-Means', '1', 'Finds behavioral segments that may be more meaningful than demographic labels'],
        ['XGBoost', '2', 'Handles non-linear relationships, mixed feature types, and produces feature importance rankings'],
        ['LightGBM', '2', 'Validates XGBoost findings through model agreement; leaf-wise growth may capture different patterns'],
        ['SHAP', '2', 'Moves from "what matters" to "how much and in which direction for each individual" — enables actionable recommendations'],
    ]
)

# ===== SAVE =====
doc.save('PROJECT_DOCUMENTATION.docx')
print('PROJECT_DOCUMENTATION.docx generated successfully.')
